from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import time
import uuid
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

import secrets

from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse
from pydantic import BaseModel
from starlette.responses import StreamingResponse
from pypdf import PdfReader
from docx import Document as DocxDocument

# Native AI Providers (v2.4.0 - fully native, no msgmodel dependency)
from app.providers.gemini import (
    GeminiProvider as NativeGeminiProvider,
    stream_gemini,
    get_gemini_privacy_info,
    GeminiAPIError,
    RateLimitError as GeminiRateLimitError,
)

from app.providers.openai import (
    OpenAIProvider as NativeOpenAIProvider,
    stream_openai,
    get_openai_privacy_info,
    OpenAIAPIError,
    RateLimitError as OpenAIRateLimitError,
)

from app.providers.anthropic import (
    AnthropicProvider as NativeAnthropicProvider,
    stream_anthropic,
    get_anthropic_privacy_info,
    AnthropicAPIError,
    RateLimitError as AnthropicRateLimitError,
)

# Key management
from app.key_storage import (
    get_configured_providers,
    save_key,
    delete_key,
    get_provider_key_for_use,
    SUPPORTED_PROVIDERS
)
from app.key_verification import verify_provider_key

# Authentication
from app.auth import (
    OAuthCallbackRequest,
    AuthResponse,
    UserInfo,
    exchange_oauth_code,
    create_jwt,
    require_auth,
    get_current_user,
    get_github_auth_url,
    get_google_auth_url,
    get_apple_auth_url,
    get_microsoft_auth_url,
    REQUIRE_AUTH,
)

# Database
from app.database import (
    init_db,
    close_db,
    create_user,
    get_user_quota,
    increment_quota,
    FREE_TIER_QUOTA,
    PAID_TIER_QUOTA,
)

# Billing
from app.billing import router as billing_router

# -----------------------------
# Logging Configuration
# -----------------------------
# Set LOG_LEVEL environment variable to control verbosity
# Production: LOG_LEVEL=WARNING (or ERROR)
# Development: LOG_LEVEL=DEBUG
log_level_env = os.environ.get("LOG_LEVEL", "WARNING").upper()
allow_debug_logs = os.environ.get("ALLOW_DEBUG_LOGS", "false").lower() == "true"
allowed_levels = {"CRITICAL", "ERROR", "WARNING", "INFO"}

# Clamp log level to privacy-safe defaults; DEBUG requires explicit opt-in
if log_level_env == "DEBUG" and not allow_debug_logs:
    log_level_env = "INFO"
elif log_level_env not in allowed_levels:
    log_level_env = "WARNING"

logging.basicConfig(
    level=getattr(logging, log_level_env, logging.WARNING),
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)
logger = logging.getLogger(__name__)

# Suppress access logs by default to avoid leaking paths/query params
if os.environ.get("DISABLE_ACCESS_LOGS", "true").lower() == "true":
    logging.getLogger("uvicorn.access").disabled = True

# API key configuration (mandatory by default)
API_KEY = os.environ.get("API_KEY", "")
REQUIRE_API_KEY = os.environ.get("REQUIRE_API_KEY", "true").lower() != "false"
API_KEY_HEADER = "x-api-key"

if REQUIRE_API_KEY and not API_KEY:
    raise RuntimeError("API_KEY must be set when REQUIRE_API_KEY is enabled (default)")

# -----------------------------
# Platform API Keys (for quota-based usage)
# -----------------------------
# These are botchat's own API keys used for users without BYOK
# Keys are stored in GCP Secret Manager and injected as env vars at deploy time
PLATFORM_OPENAI_KEY = os.environ.get("PLATFORM_OPENAI_API_KEY", "")
PLATFORM_GEMINI_KEY = os.environ.get("PLATFORM_GEMINI_API_KEY", "")
PLATFORM_ANTHROPIC_KEY = os.environ.get("PLATFORM_ANTHROPIC_API_KEY", "")

PLATFORM_KEYS = {
    "openai": PLATFORM_OPENAI_KEY,
    "gemini": PLATFORM_GEMINI_KEY,
    "anthropic": PLATFORM_ANTHROPIC_KEY,
}


def get_platform_key(provider: str) -> Optional[str]:
    """Get platform API key for a provider if available."""
    return PLATFORM_KEYS.get(provider.lower()) or None


def has_platform_key(provider: str) -> bool:
    """Check if platform has a key configured for this provider."""
    return bool(get_platform_key(provider))


# -----------------------------
# PDF Text Extraction Utilities
# -----------------------------

def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from PDF file.
    
    Returns extracted text, or empty string if extraction fails.
    Note: This may return empty string for scanned/image-based PDFs.
    """
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text and text.strip():  # Only add if text is non-empty after stripping
                    # Add page indicator for clarity
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            except Exception as page_error:
                logger.warning("Failed to extract text from page %d: %s", page_num + 1, str(page_error))
                continue
        
        extracted = "\n\n".join(text_parts)
        return extracted
    except Exception as e:
        logger.warning("PDF text extraction failed: %s", str(e))
        return ""


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from DOCX file.
    
    Returns extracted text, or empty string if extraction fails.
    """
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = DocxDocument(docx_file)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        
        extracted = "\n\n".join(text_parts)
        return extracted
    except Exception as e:
        logger.warning("DOCX text extraction failed: %s", str(e))
        return ""


def prepare_file_for_provider(
    file_bytes: bytes, 
    filename: str, 
    provider: str
) -> Tuple[Optional[io.BytesIO], Optional[str], Optional[str]]:
    """Prepare file for submission to specific provider.
    
    Returns (file_like, filename, extracted_text_for_message)
    
    For PDF files:
    - Gemini: Returns original PDF (native support)
    - OpenAI/Anthropic: Extracts text and returns it as a text attachment metadata
              If extraction fails, includes a warning message about the PDF
    
    For DOCX files:
    - All providers: Extract text (no provider supports DOCX natively)
    
    For other files: Returns original file for all providers
    """
    filename_lower = filename.lower()
    is_pdf = filename_lower.endswith('.pdf')
    is_docx = filename_lower.endswith('.docx')
    is_legacy_doc = filename_lower.endswith('.doc') and not filename_lower.endswith('.docx')
    
    # Legacy .DOC format - not supported, ask user to convert
    if is_legacy_doc:
        warning_msg = f"[DOC Error] The legacy .doc format is not supported. Please save '{filename}' as .docx (Word 2007+) and try again."
        logger.warning("%s", warning_msg)
        return None, None, warning_msg
    
    # DOCX handling - extract text for all providers (none support DOCX natively)
    if is_docx:
        extracted_text = extract_docx_text(file_bytes)
        if extracted_text and extracted_text.strip():
            logger.info("DOCX text extracted from %s: %d chars", filename, len(extracted_text))
            return None, None, extracted_text
        else:
            warning_msg = f"[DOCX Error] Could not extract text from '{filename}'. The document may be empty, password-protected, or in an unsupported format."
            logger.warning("%s", warning_msg)
            return None, None, warning_msg
    
    # PDF handling
    if is_pdf:
        if provider.lower() == "gemini":
            # Gemini supports PDFs natively
            file_like = io.BytesIO(file_bytes)
            file_like.seek(0)
            return file_like, filename, None
        elif provider.lower() in ("openai", "anthropic"):
            # OpenAI and Anthropic: Extract text from PDF
            # Neither supports native PDF - Anthropic has strict request size limits
            extracted_text = extract_pdf_text(file_bytes)
            if extracted_text and extracted_text.strip():
                logger.debug("PDF text extracted from %s: %d chars", filename, len(extracted_text))
                # Return None for file_like since we're using text instead
                return None, None, extracted_text
            else:
                # PDF extraction failed or returned empty content
                # This typically happens with scanned/image-based PDFs
                warning_msg = f"[PDF Error] Could not extract text from '{filename}'. This is typically because the PDF contains scanned images or OCR-protected content. Please provide a text-based PDF or paste the content as text."
                logger.warning("%s", warning_msg)
                return None, None, warning_msg
        else:
            # Other providers: return original PDF
            file_like = io.BytesIO(file_bytes)
            file_like.seek(0)
            return file_like, filename, None
    
    # Non-PDF/DOCX files: pass through unchanged
    file_like = io.BytesIO(file_bytes)
    file_like.seek(0)
    return file_like, filename, None


def get_privacy_info(provider: str, api_key: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """Get privacy metadata for a provider.
    
    Args:
        provider: Provider name (openai, gemini, anthropic)
        api_key: For determining BYOK vs platform mode privacy info
    
    Returns privacy info dict or None if provider is unknown.
    """
    provider_lower = provider.lower()
    
    try:
        if provider_lower == "openai":
            return get_openai_privacy_info(api_key=api_key)
        elif provider_lower == "gemini":
            return get_gemini_privacy_info(api_key=api_key)
        elif provider_lower == "anthropic":
            return get_anthropic_privacy_info(api_key=api_key)
        else:
            logger.warning("Unknown provider for privacy info: %s", provider)
            return None
            return None
    except Exception as e:
        logger.error("Failed to fetch privacy info for %s: %s", provider, str(e))
        return None

# -----------------------------
# Models
# -----------------------------

class ModelConfig(BaseModel):
    id: str  # client-side id; stable key for UI panel
    provider: str
    model: str
    system: str = ""
    max_tokens: Optional[int] = None  # Optional per-persona token limit
    provider_key: Optional[str] = None  # Client-provided API key (decrypted from localStorage)
    web_search_enabled: bool = False  # Enable web search tool (model decides when to search)


class RunCreateRequest(BaseModel):
    message: str
    configs: List[ModelConfig]
    max_parallel: int = 10


class RunCreateResponse(BaseModel):
    run_id: str


class SynthesizeRequest(BaseModel):
    system: str = "You synthesize multiple model answers into one best answer."
    instruction: str = "Synthesize the answers."
    include_config_ids: Optional[List[str]] = None


# Settings Models
class SaveKeyRequest(BaseModel):
    provider: str
    api_key: str


class VerifyKeyRequest(BaseModel):
    provider: str
    api_key: str


# -----------------------------
# In-memory run state (dev-only)
# -----------------------------

@dataclass
class RunState:
    run_id: str
    created_at: float
    queue: asyncio.Queue[str] = field(default_factory=asyncio.Queue)
    finals: Dict[str, str] = field(default_factory=dict)  # config_id -> final text
    done: bool = False
    # Quota tracking for platform key usage
    user_info: Optional[UserInfo] = None  # User context for quota updates
    platform_key_configs: List[str] = field(default_factory=list)  # config_ids using platform keys
    successful_platform_configs: List[str] = field(default_factory=list)  # track successes for quota


RUNS: Dict[str, RunState] = {}
RUNS_LOCK = asyncio.Lock()

# Privacy: Auto-cleanup stale runs to prevent memory accumulation of sensitive data
RUN_TTL_SECONDS = max(10, int(os.environ.get("RUN_TTL_SECONDS", "60")))


async def require_api_key_or_jwt(
    x_api_key: Optional[str] = Header(None, alias=API_KEY_HEADER),
    authorization: Optional[str] = Header(None),
) -> Optional[UserInfo]:
    """Enforce authentication via API key OR JWT token.
    
    Supports both legacy API key auth and new JWT auth:
    - x-api-key header: Legacy static key (for dev/self-hosted)
    - Authorization: Bearer <token>: JWT from OAuth flow
    
    Returns UserInfo if JWT auth, None if API key auth.
    """
    # If auth is disabled entirely, allow through
    if not REQUIRE_API_KEY and not REQUIRE_AUTH:
        return None
    
    # Try JWT first (new auth flow)
    if authorization and authorization.startswith("Bearer "):
        from app.auth import verify_jwt
        token = authorization[7:]
        return verify_jwt(token)
    
    # Fall back to API key (legacy/self-hosted)
    if REQUIRE_API_KEY:
        if x_api_key and API_KEY and secrets.compare_digest(x_api_key, API_KEY):
            return None  # API key valid, no user info
        
    # If we got here with REQUIRE_AUTH enabled, we need a valid JWT
    if REQUIRE_AUTH:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    # REQUIRE_API_KEY is true but key is invalid
    if REQUIRE_API_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized")
    
    return None


# Alias for backward compatibility
async def require_api_key(
    x_api_key: Optional[str] = Header(None, alias=API_KEY_HEADER),
    authorization: Optional[str] = Header(None),
) -> None:
    """Legacy wrapper - now supports both API key and JWT."""
    await require_api_key_or_jwt(x_api_key, authorization)


async def cleanup_stale_runs() -> None:
    """Remove completed runs older than TTL to prevent memory accumulation."""
    while True:
        await asyncio.sleep(60)  # Check every minute
        now = time.time()
        async with RUNS_LOCK:
            stale_ids = [
                run_id for run_id, run in RUNS.items()
                if run.done and (now - run.created_at) > RUN_TTL_SECONDS
            ]
            for run_id in stale_ids:
                del RUNS[run_id]
                logger.debug("Cleaned up stale run: %s", run_id)


# -----------------------------
# FastAPI setup
# -----------------------------

@asynccontextmanager
async def lifespan(app: FastAPI):
    """FastAPI lifespan context manager.
    
    Handles startup and shutdown events.
    Replaces deprecated @app.on_event("startup") pattern.
    """
    # Startup
    await init_db()  # Initialize database connection
    asyncio.create_task(cleanup_stale_runs())
    yield
    # Shutdown
    await close_db()  # Close database connection


app = FastAPI(title="botchat-backend", lifespan=lifespan)

# Include billing routes
app.include_router(billing_router)


@app.get("/health")
async def health_check():
    """Health check endpoint for Docker and monitoring."""
    from app.version import __version__
    return {"status": "healthy", "service": "botchat-backend", "version": __version__}


# CORS Configuration
# Set CORS_ORIGINS environment variable for production (comma-separated)
# Example: CORS_ORIGINS=https://myapp.example.com,https://app.example.com
# Default includes development origins plus production Cloud Run and custom domain
default_cors = "http://localhost:3000,http://localhost:3001,http://localhost:5173,https://botchat-frontend-887036129720.us-central1.run.app,https://dev.botchat.ca"
cors_origins_env = os.environ.get("CORS_ORIGINS", default_cors)

if "*" in cors_origins_env:
    raise RuntimeError("CORS_ORIGINS must not contain '*' to maintain origin pinning")

cors_origins = [origin.strip() for origin in cors_origins_env.split(",") if origin.strip()]

if not cors_origins:
    raise RuntimeError("CORS_ORIGINS must specify at least one allowed origin")

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,  # Allow cookies/auth for OAuth flow
    allow_methods=["GET", "POST", "DELETE", "OPTIONS", "HEAD"],
    allow_headers=["Content-Type", API_KEY_HEADER, "Authorization", "Accept", "Origin"],
    expose_headers=["Content-Type"],
    max_age=3600,  # Cache preflight for 1 hour
)


# -----------------------------
# SSE helpers
# -----------------------------

def sse(event: str, data: Any) -> str:
    # EventSource expects lines: "event: X\n" + "data: ...\n\n"
    payload = json.dumps(data, ensure_ascii=False)
    return f"event: {event}\ndata: {payload}\n\n"


def _get_mime_type(filename: str) -> str:
    """Get MIME type from filename extension.
    
    Used for file attachments sent to AI providers.
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    mime_types = {
        # Documents
        'pdf': 'application/pdf',
        'doc': 'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain',
        'md': 'text/markdown',
        'rtf': 'application/rtf',
        # Data formats
        'json': 'application/json',
        'csv': 'text/csv',
        'xml': 'application/xml',
        'yaml': 'application/yaml',
        'yml': 'application/yaml',
        # Code files
        'py': 'text/x-python',
        'js': 'text/javascript',
        'ts': 'text/typescript',
        'html': 'text/html',
        'css': 'text/css',
        # Images
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'gif': 'image/gif',
        'webp': 'image/webp',
        'heic': 'image/heic',
    }
    result = mime_types.get(ext, 'application/octet-stream')
    logger.info("_get_mime_type: filename='%s', ext='%s', mime_type='%s'", filename, ext, result)
    return result


async def publish(run: RunState, event: str, data: Any) -> None:
    await run.queue.put(sse(event, data))


# -----------------------------
# Provider Streaming
# -----------------------------

async def stream_to_provider(run: RunState, cfg: ModelConfig, message: str, file_like: Optional[io.BytesIO] = None, filename: Optional[str] = None, extracted_text: Optional[str] = None) -> None:
    """Stream one provider's response to a message (optionally with attachment).
    
    v2.4.0 ARCHITECTURE - All native providers:
    - Gemini: Native provider via Vertex AI (platform) or AI Studio (BYOK)
    - OpenAI: Native provider (platform or BYOK)
    - Anthropic: Native provider (platform or BYOK)
    
    For PDFs with OpenAI/Anthropic: extracted_text contains the PDF text content.
    
    Privacy metadata is fetched and sent to the frontend for transparency.
    Each request is completely independent and stateless.
    """
    await publish(run, "panel_start", {"config_id": cfg.id})

    # Fetch and send privacy metadata for this provider
    # Pass api_key to determine BYOK vs platform privacy info
    privacy_info = get_privacy_info(cfg.provider, api_key=cfg.provider_key)
    await publish(run, "panel_privacy", {
        "config_id": cfg.id,
        "privacy": privacy_info
    })

    loop = asyncio.get_running_loop()
    final_parts: list[str] = []
    
    # Route to appropriate native streaming implementation
    provider_lower = cfg.provider.lower()
    
    if provider_lower == "gemini":
        await _stream_gemini_native(run, cfg, message, file_like, filename, extracted_text, loop, final_parts)
    elif provider_lower == "openai":
        await _stream_openai_native(run, cfg, message, file_like, filename, extracted_text, loop, final_parts)
    elif provider_lower == "anthropic":
        await _stream_anthropic_native(run, cfg, message, file_like, filename, extracted_text, loop, final_parts)
    else:
        # Unknown provider
        await publish(run, "panel_error", {"config_id": cfg.id, "error": f"Unknown provider: {cfg.provider}"})


async def _stream_gemini_native(
    run: RunState, 
    cfg: ModelConfig, 
    message: str, 
    file_like: Optional[io.BytesIO], 
    filename: Optional[str], 
    extracted_text: Optional[str],
    loop: asyncio.AbstractEventLoop,
    final_parts: list[str]
) -> None:
    """Stream response using native Gemini provider (Vertex AI or AI Studio).
    
    v2.2.0: Replaces msgmodel for Gemini with direct SDK integration.
    - Platform usage: Vertex AI with service account
    - BYOK usage: AI Studio with user's API key
    """
    def _run_gemini_thread():
        max_retries = 3
        retry_count = 0
        max_tokens_to_use = cfg.max_tokens or 4000
        
        # Determine which key to use (None = platform/Vertex AI)
        api_key = cfg.provider_key  # None for platform, user key for BYOK
        
        while retry_count <= max_retries:
            try:
                # Build effective message with any extracted text
                effective_message = message
                if extracted_text:
                    effective_message = f"{extracted_text}\n\n---\n\n[User Query]\n{message}"
                    logger.debug("Using extracted text (%d chars) for Gemini", len(extracted_text))
                
                # Build file data for native provider
                file_data = None
                if file_like and filename:
                    file_like.seek(0)
                    file_bytes = file_like.read()
                    # Determine MIME type from filename
                    mime_type = _get_mime_type(filename)
                    file_data = [{"bytes": file_bytes, "mime_type": mime_type, "name": filename}]
                    logger.debug("File prepared for Gemini: %s (%s, %d bytes)", 
                               filename, mime_type, len(file_bytes))
                
                logger.debug("Calling native Gemini provider: model=%s, backend=%s (attempt %d/%d)",
                           cfg.model, "ai_studio" if api_key else "vertex_ai", 
                           retry_count + 1, max_retries + 1)
                
                final_parts.clear()
                result = None
                gen = stream_gemini(
                    message=effective_message,
                    model=cfg.model,
                    api_key=api_key,
                    system_instruction=cfg.system if cfg.system else None,
                    max_tokens=max_tokens_to_use,
                    file_data=file_data,
                    web_search_enabled=cfg.web_search_enabled,
                )
                
                # Iterate through generator, capturing return value
                try:
                    while True:
                        chunk = next(gen)
                        final_parts.append(chunk)
                        loop.call_soon_threadsafe(
                            run.queue.put_nowait,
                            sse("panel_token", {"config_id": cfg.id, "token": chunk}),
                        )
                except StopIteration as e:
                    result = e.value  # Capture generator return value (citations dict)
                
                final = "".join(final_parts).strip()
                run.finals[cfg.id] = final
                
                if cfg.id in run.platform_key_configs:
                    run.successful_platform_configs.append(cfg.id)
                
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_final", {"config_id": cfg.id, "final": final}),
                )
                
                # Emit citations if web search was used
                if result and result.get('citations'):
                    loop.call_soon_threadsafe(
                        run.queue.put_nowait,
                        sse("panel_citations", {"config_id": cfg.id, "citations": result['citations']}),
                    )
                    logger.debug("Emitted %d citations for Gemini panel %s", 
                               len(result['citations']), cfg.id)
                
                return  # Success
                
            except GeminiRateLimitError as e:
                if retry_count < max_retries:
                    wait_time = 2 ** retry_count
                    retry_count += 1
                    logger.warning("Rate limited on Gemini. Retrying in %ds (attempt %d/%d). Error: %s",
                                 wait_time, retry_count, max_retries, str(e))
                    time.sleep(wait_time)
                    continue
                else:
                    error_msg = f"Rate limited after {retry_count} retries: {str(e)}"
                    logger.error("Gemini rate limit exceeded: %s", error_msg)
                    loop.call_soon_threadsafe(
                        run.queue.put_nowait,
                        sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                    )
                    return
                    
            except GeminiAPIError as e:
                error_msg = str(e)
                logger.error("Gemini API error: %s", error_msg, exc_info=True)
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                )
                return
                
            except Exception as e:
                error_msg = f"{type(e).__name__}: {str(e)}"
                logger.error("Unexpected error in Gemini stream: %s", error_msg, exc_info=True)
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                )
                return
    
    try:
        await asyncio.to_thread(_run_gemini_thread)
        final = run.finals.get(cfg.id, "")
        if not final:
            await publish(run, "panel_final", {"config_id": cfg.id, "final": final})
    except Exception as e:
        await publish(run, "panel_error", {"config_id": cfg.id, "error": str(e)})


async def _stream_openai_native(
    run: RunState, 
    cfg: ModelConfig, 
    message: str, 
    file_like: Optional[io.BytesIO], 
    filename: Optional[str], 
    extracted_text: Optional[str],
    loop: asyncio.AbstractEventLoop,
    final_parts: list[str]
) -> None:
    """Stream response using native OpenAI provider.
    
    v2.3.0: Replaces msgmodel for OpenAI with direct SDK integration.
    - Platform usage: botchat's API key (no ZDR agreement)
    - BYOK usage: User's API key (ZDR honored if they have enterprise agreement)
    
    Privacy note: X-OpenAI-No-Store header is always sent, but only honored
    for organizations with formal ZDR agreements.
    """
    def _run_openai_thread():
        max_retries = 3
        retry_count = 0
        max_tokens_to_use = cfg.max_tokens or 4000
        
        # Determine which key to use (None = platform key)
        api_key = cfg.provider_key  # None for platform, user key for BYOK
        
        while retry_count <= max_retries:
            try:
                # Build effective message with any extracted text (PDFs)
                effective_message = message
                if extracted_text:
                    effective_message = f"{extracted_text}\n\n---\n\n[User Query]\n{message}"
                    logger.debug("Using extracted text (%d chars) for OpenAI", len(extracted_text))
                
                # Build file data for native provider (images only)
                file_data = None
                logger.info("OpenAI file check: file_like=%s, filename=%s", file_like is not None, filename)
                if file_like and filename:
                    file_like.seek(0)
                    file_bytes = file_like.read()
                    mime_type = _get_mime_type(filename)
                    logger.info("OpenAI: file_bytes=%d, mime_type=%s", len(file_bytes), mime_type)
                    
                    # OpenAI only supports images natively, PDFs are text-extracted
                    if mime_type.startswith("image/"):
                        file_data = [{"bytes": file_bytes, "mime_type": mime_type, "name": filename}]
                        logger.info("Image prepared for OpenAI: %s (%s, %d bytes)", 
                                   filename, mime_type, len(file_bytes))
                    else:
                        logger.debug("Non-image file %s handled via text extraction", filename)
                
                logger.debug("Calling native OpenAI provider: model=%s, backend=%s (attempt %d/%d)",
                           cfg.model, "byok" if api_key else "platform", 
                           retry_count + 1, max_retries + 1)
                
                final_parts.clear()
                result = None
                gen = stream_openai(
                    message=effective_message,
                    model=cfg.model,
                    api_key=api_key,
                    system_instruction=cfg.system if cfg.system else None,
                    max_tokens=max_tokens_to_use,
                    file_data=file_data,
                    web_search_enabled=cfg.web_search_enabled,
                )
                
                # Iterate through generator, capturing return value
                try:
                    while True:
                        chunk = next(gen)
                        final_parts.append(chunk)
                        loop.call_soon_threadsafe(
                            run.queue.put_nowait,
                            sse("panel_token", {"config_id": cfg.id, "token": chunk}),
                        )
                except StopIteration as e:
                    result = e.value  # Capture generator return value (citations dict)
                
                final = "".join(final_parts).strip()
                run.finals[cfg.id] = final
                
                if cfg.id in run.platform_key_configs:
                    run.successful_platform_configs.append(cfg.id)
                
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_final", {"config_id": cfg.id, "final": final}),
                )
                
                # Emit citations if web search was used
                if result and result.get('citations'):
                    loop.call_soon_threadsafe(
                        run.queue.put_nowait,
                        sse("panel_citations", {"config_id": cfg.id, "citations": result['citations']}),
                    )
                    logger.debug("Emitted %d citations for OpenAI panel %s", 
                               len(result['citations']), cfg.id)
                
                return  # Success
                
            except OpenAIRateLimitError as e:
                if retry_count < max_retries:
                    wait_time = 2 ** retry_count
                    retry_count += 1
                    logger.warning("Rate limited on OpenAI. Retrying in %ds (attempt %d/%d). Error: %s",
                                 wait_time, retry_count, max_retries, str(e))
                    time.sleep(wait_time)
                    continue
                else:
                    error_msg = f"Rate limited after {retry_count} retries: {str(e)}"
                    logger.error("OpenAI rate limit exceeded: %s", error_msg)
                    loop.call_soon_threadsafe(
                        run.queue.put_nowait,
                        sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                    )
                    return
                    
            except OpenAIAPIError as e:
                error_msg = str(e)
                logger.error("OpenAI API error: %s", error_msg, exc_info=True)
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                )
                return
                
            except Exception as e:
                error_msg = f"{type(e).__name__}: {str(e)}"
                logger.error("Unexpected error in OpenAI stream: %s", error_msg, exc_info=True)
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                )
                return
    
    try:
        await asyncio.to_thread(_run_openai_thread)
        final = run.finals.get(cfg.id, "")
        if not final:
            await publish(run, "panel_final", {"config_id": cfg.id, "final": final})
    except Exception as e:
        await publish(run, "panel_error", {"config_id": cfg.id, "error": str(e)})


async def _stream_anthropic_native(
    run: RunState, 
    cfg: ModelConfig, 
    message: str, 
    file_like: Optional[io.BytesIO], 
    filename: Optional[str], 
    extracted_text: Optional[str],
    loop: asyncio.AbstractEventLoop,
    final_parts: list[str]
) -> None:
    """Stream response using native Anthropic provider.
    
    v2.4.0: Replaces msgmodel for Anthropic with direct SDK integration.
    - Platform usage: botchat's API key
    - BYOK usage: User's API key
    
    Privacy note: Anthropic does NOT use API data for training by default.
    Data may be retained up to 30 days for trust & safety monitoring.
    """
    def _run_anthropic_thread():
        max_retries = 3
        retry_count = 0
        max_tokens_to_use = cfg.max_tokens or 4096
        
        # Determine which key to use (None = platform key)
        api_key = cfg.provider_key  # None for platform, user key for BYOK
        
        while retry_count <= max_retries:
            try:
                # Build effective message with any extracted text (PDFs)
                effective_message = message
                if extracted_text:
                    effective_message = f"{extracted_text}\n\n---\n\n[User Query]\n{message}"
                    logger.debug("Using extracted text (%d chars) for Anthropic", len(extracted_text))
                
                # Build file data for native provider (images only)
                file_data = None
                logger.info("Anthropic file check: file_like=%s, filename=%s", file_like is not None, filename)
                if file_like and filename:
                    file_like.seek(0)
                    file_bytes = file_like.read()
                    mime_type = _get_mime_type(filename)
                    logger.info("Anthropic: file_bytes=%d, mime_type=%s", len(file_bytes), mime_type)
                    
                    # Anthropic only supports images natively, PDFs are text-extracted
                    if mime_type.startswith("image/"):
                        file_data = [{"bytes": file_bytes, "mime_type": mime_type, "name": filename}]
                        logger.info("Image prepared for Anthropic: %s (%s, %d bytes)", 
                                   filename, mime_type, len(file_bytes))
                    else:
                        logger.debug("Non-image file %s handled via text extraction", filename)
                
                logger.debug("Calling native Anthropic provider: model=%s, backend=%s (attempt %d/%d)",
                           cfg.model, "byok" if api_key else "platform", 
                           retry_count + 1, max_retries + 1)
                
                final_parts.clear()
                result = None
                gen = stream_anthropic(
                    message=effective_message,
                    model=cfg.model,
                    api_key=api_key,
                    system_instruction=cfg.system if cfg.system else None,
                    max_tokens=max_tokens_to_use,
                    file_data=file_data,
                    web_search_enabled=cfg.web_search_enabled,
                )
                
                # Iterate through generator, capturing return value
                try:
                    while True:
                        chunk = next(gen)
                        final_parts.append(chunk)
                        loop.call_soon_threadsafe(
                            run.queue.put_nowait,
                            sse("panel_token", {"config_id": cfg.id, "token": chunk}),
                        )
                except StopIteration as e:
                    result = e.value  # Capture generator return value (citations dict)
                
                final = "".join(final_parts).strip()
                run.finals[cfg.id] = final
                
                if cfg.id in run.platform_key_configs:
                    run.successful_platform_configs.append(cfg.id)
                
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_final", {"config_id": cfg.id, "final": final}),
                )
                
                # Emit citations if web search was used
                if result and result.get('citations'):
                    loop.call_soon_threadsafe(
                        run.queue.put_nowait,
                        sse("panel_citations", {"config_id": cfg.id, "citations": result['citations']}),
                    )
                    logger.debug("Emitted %d citations for Anthropic panel %s", 
                               len(result['citations']), cfg.id)
                
                return  # Success
                
            except AnthropicRateLimitError as e:
                if retry_count < max_retries:
                    wait_time = 2 ** retry_count
                    retry_count += 1
                    logger.warning("Rate limited on Anthropic. Retrying in %ds (attempt %d/%d). Error: %s",
                                 wait_time, retry_count, max_retries, str(e))
                    time.sleep(wait_time)
                    continue
                else:
                    error_msg = f"Rate limited after {retry_count} retries: {str(e)}"
                    logger.error("Anthropic rate limit exceeded: %s", error_msg)
                    loop.call_soon_threadsafe(
                        run.queue.put_nowait,
                        sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                    )
                    return
                    
            except AnthropicAPIError as e:
                error_msg = str(e)
                logger.error("Anthropic API error: %s", error_msg, exc_info=True)
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                )
                return
                
            except Exception as e:
                error_msg = f"{type(e).__name__}: {str(e)}"
                logger.error("Unexpected error in Anthropic stream: %s", error_msg, exc_info=True)
                loop.call_soon_threadsafe(
                    run.queue.put_nowait,
                    sse("panel_error", {"config_id": cfg.id, "error": error_msg}),
                )
                return
    
    try:
        await asyncio.to_thread(_run_anthropic_thread)
        final = run.finals.get(cfg.id, "")
        if not final:
            await publish(run, "panel_final", {"config_id": cfg.id, "final": final})
    except Exception as e:
        await publish(run, "panel_error", {"config_id": cfg.id, "error": str(e)})


async def run_fanout(run: RunState, configs: List[ModelConfig], message: str, max_parallel: int, file_list: Optional[List[Tuple[bytes, str]]] = None) -> None:
    try:
        await publish(run, "run_start", {"run_id": run.run_id, "n": len(configs)})

        # Limit concurrency
        sem = asyncio.Semaphore(max_parallel)

        async def guarded(cfg: ModelConfig):
            async with sem:
                # For multiple files: extract text from all and send as text
                # For single file: use native format when possible
                file_like = None
                file_name = None
                extracted_text = None
                
                if file_list:
                    if len(file_list) > 1:
                        # Multiple files: Extract text from all files and combine
                        # This ensures all files are visible to the provider
                        logger.debug("Multiple files detected (%d), extracting text from all", len(file_list))
                        all_texts = []
                        for file_bytes, filename in file_list:
                            is_pdf = filename.lower().endswith('.pdf')
                            if is_pdf:
                                # Extract text from PDF
                                pdf_text = extract_pdf_text(file_bytes)
                                if pdf_text and pdf_text.strip():
                                    all_texts.append(f"\n[File: {filename}]\n{pdf_text}")
                                else:
                                    all_texts.append(f"\n[File: {filename}] (Failed to extract text - may be scanned/image PDF)")
                            else:
                                # For non-PDF files, read as text
                                try:
                                    text = file_bytes.decode('utf-8', errors='replace')
                                    all_texts.append(f"\n[File: {filename}]\n{text}")
                                except Exception as e:
                                    all_texts.append(f"\n[File: {filename}] (Could not read file: {str(e)})")
                        
                        if all_texts:
                            extracted_text = "".join(all_texts)
                            logger.debug("Combined text from %d files: %d chars", len(file_list), len(extracted_text))
                    else:
                        # Single file: use prepare_file_for_provider for native format support
                        file_bytes, filename = file_list[0]
                        file_like, file_name, extracted_text = prepare_file_for_provider(
                            file_bytes,
                            filename,
                            cfg.provider
                        )
                        logger.debug("Single file: %s, native_support=%s, extracted_text=%s", filename, file_like is not None, extracted_text is not None)
                
                await stream_to_provider(run, cfg, message, file_like, file_name, extracted_text)

        tasks = [asyncio.create_task(guarded(cfg)) for cfg in configs]
        await asyncio.gather(*tasks)

        # Increment quota BEFORE sending run_done so frontend sees updated quota
        quota_updated = None
        if run.user_info and run.successful_platform_configs:
            success_count = len(run.successful_platform_configs)
            try:
                quota_updated = await increment_quota(
                    run.user_info.provider,
                    run.user_info.user_id,
                    count=success_count
                )
                if quota_updated:
                    logger.debug(
                        "Quota incremented: user=%s:%s, count=%d, new_used=%d/%d",
                        run.user_info.provider,
                        run.user_info.user_id,
                        success_count,
                        quota_updated["used"],
                        quota_updated["limit"]
                    )
            except Exception as e:
                logger.error("Failed to increment quota: %s", str(e))

        # Include quota info in run_done event for immediate UI update
        run_done_data = {"run_id": run.run_id}
        if quota_updated:
            run_done_data["quota"] = quota_updated
        await publish(run, "run_done", run_done_data)
    except Exception as e:
        await publish(run, "run_error", {"run_id": run.run_id, "error": str(e)})
    finally:
        run.done = True


async def fake_synthesize(run: RunState, synth_req: SynthesizeRequest) -> None:
    # Very simple fake synth for now
    include = synth_req.include_config_ids or list(run.finals.keys())
    parts = [run.finals[cid] for cid in include if cid in run.finals]

    await publish(run, "synth_start", {"include": include})

    # Stream a pretend synthesis
    combined = " / ".join(parts)
    words = (synth_req.instruction + " " + combined).split(" ")
    buf = []
    for w in words:
        buf.append(w)
        await publish(run, "synth_token", {"token": w + " "})
        await asyncio.sleep(0.02)

    await publish(run, "synth_final", {"final": "".join(buf).strip()})


# -----------------------------
# Routes: Authentication
# -----------------------------

class AuthUrlRequest(BaseModel):
    """Request for OAuth authorization URL."""
    provider: str  # "github" or "google"
    redirect_uri: str


class AuthUrlResponse(BaseModel):
    """OAuth authorization URL response."""
    url: str


@app.post("/auth/url", response_model=AuthUrlResponse)
async def get_auth_url(req: AuthUrlRequest):
    """Get OAuth authorization URL for a provider.
    
    Frontend redirects user to this URL to start OAuth flow.
    
    Note: For Apple, we use the backend callback URL because Apple requires
    form_post response mode which sends a POST request. The backend receives
    the POST and redirects to the frontend with the code as a query param.
    """
    if req.provider == "github":
        url = get_github_auth_url(req.redirect_uri)
    elif req.provider == "google":
        url = get_google_auth_url(req.redirect_uri)
    elif req.provider == "apple":
        # Apple needs to POST to backend, which then redirects to frontend
        # BACKEND_URL must be set in environment for Apple Sign In to work
        backend_url = os.environ.get("BACKEND_URL", "")
        if not backend_url:
            raise HTTPException(status_code=500, detail="BACKEND_URL not configured for Apple Sign In")
        apple_redirect = f"{backend_url}/auth/apple/callback"
        url = get_apple_auth_url(apple_redirect)
    elif req.provider == "microsoft":
        url = get_microsoft_auth_url(req.redirect_uri)
    else:
        raise HTTPException(status_code=400, detail=f"Unknown provider: {req.provider}")
    
    return AuthUrlResponse(url=url)


# Frontend URL for redirects (needed for Apple callback)
FRONTEND_URL = os.environ.get("FRONTEND_URL", "http://localhost:5173")


@app.post("/auth/apple/callback")
async def apple_callback_post(
    code: str = Form(...),
    state: Optional[str] = Form(None),
    id_token: Optional[str] = Form(None),
    user: Optional[str] = Form(None),  # Apple sends user info as JSON string on first auth
):
    """Handle Apple's form POST callback.
    
    Apple Sign In uses response_mode=form_post, which sends the authorization
    code via POST. This endpoint receives that POST and redirects to the
    frontend callback page with the code as a query parameter.
    """
    # Build redirect URL to frontend
    redirect_url = f"{FRONTEND_URL}/auth/callback?code={code}&provider=apple"
    if state:
        redirect_url += f"&state={state}"
    
    return RedirectResponse(url=redirect_url, status_code=302)


@app.post("/auth/callback", response_model=AuthResponse)
async def auth_callback(req: OAuthCallbackRequest):
    """Exchange OAuth code for JWT.
    
    Called by frontend after user completes OAuth flow.
    Creates/links user account by email and returns JWT token.
    
    Email linking: If a user with the same email already exists (from a
    different OAuth provider), they are linked to the same account. This
    ensures subscriptions work across login methods.
    """
    from app.auth import check_email_allowed
    
    # Get user info from OAuth provider
    user = await exchange_oauth_code(req)
    
    # Check email allowlist (if configured)
    check_email_allowed(user.email)
    
    # Create or link user account in database (links by email if exists)
    try:
        db_user = await create_user(
            provider=user.provider,
            oauth_id=user.user_id,
            email=user.email,
        )
        logger.info("User authenticated: %s (db_id=%s)", user.email, db_user.get('id'))
    except Exception as e:
        # Log but don't fail - user can still authenticate
        logger.warning("Failed to persist user to database: %s", str(e))
    
    token, expires_at = create_jwt(user)
    
    return AuthResponse(
        token=token,
        user={
            "provider": user.provider,
            "id": user.user_id,
            "email": user.email,
            "name": user.name,
            "avatar": user.avatar_url,
        },
        expires_at=expires_at,
    )


@app.get("/auth/me")
async def get_current_user_info(user: UserInfo = Depends(require_auth)):
    """Get current authenticated user info.
    
    Useful for validating token and displaying user profile.
    """
    return {
        "provider": user.provider,
        "id": user.user_id,
        "email": user.email,
        "name": user.name,
        "avatar": user.avatar_url,
    }


@app.get("/auth/encryption-key")
async def get_encryption_key(user: UserInfo = Depends(require_auth)):
    """Get the user's encryption key for client-side API key encryption.
    
    This key is used to encrypt/decrypt AI provider API keys in localStorage.
    The actual API keys never touch the server - only the encryption key.
    """
    from app.database import get_user_by_oauth
    
    # Get user from database by OAuth provider and ID
    db_user = await get_user_by_oauth(user.provider, user.user_id)
    
    if not db_user or not db_user.get('encryption_key'):
        raise HTTPException(status_code=404, detail="User encryption key not found")
    
    return {
        "encryption_key": db_user['encryption_key']
    }


@app.get("/auth/quota")
async def get_quota(user: UserInfo = Depends(require_auth)):
    """Get user's message quota status.
    
    Returns:
    - used: messages used this period
    - limit: total allowed for this period (100 free, 5000 paid)
    - remaining: messages remaining
    - period_ends_at: when current period ends
    - is_paid: whether user has paid subscription
    """
    quota = await get_user_quota(user.provider, user.user_id, user.email)
    return quota


# -----------------------------
# Routes: Settings (API Key Management)
# -----------------------------

@app.get("/settings/providers")
async def get_providers(api_key=Depends(require_api_key)):
    """Get list of supported providers and their configuration status."""
    return get_configured_providers(API_KEY)


@app.post("/settings/keys/verify")
async def verify_key(req: VerifyKeyRequest, api_key=Depends(require_api_key)):
    """Verify an API key is valid without saving it."""
    if req.provider.lower() not in SUPPORTED_PROVIDERS:
        raise HTTPException(status_code=400, detail=f"Unknown provider: {req.provider}")
    
    is_valid, message, details = await verify_provider_key(req.provider, req.api_key)
    return {
        "valid": is_valid,
        "message": message,
        "details": details
    }


@app.post("/settings/keys")
async def save_provider_key(req: SaveKeyRequest, api_key=Depends(require_api_key)):
    """Save a provider API key (encrypted storage)."""
    if req.provider.lower() not in SUPPORTED_PROVIDERS:
        raise HTTPException(status_code=400, detail=f"Unknown provider: {req.provider}")
    
    # Optionally verify before saving
    is_valid, message, _ = await verify_provider_key(req.provider, req.api_key)
    if not is_valid:
        raise HTTPException(status_code=400, detail=f"Invalid API key: {message}")
    
    save_key(API_KEY, req.provider, req.api_key)
    return {"ok": True, "message": f"{req.provider} key saved successfully"}


@app.delete("/settings/keys/{provider}")
async def delete_provider_key(provider: str, api_key=Depends(require_api_key)):
    """Delete a stored provider API key."""
    if provider.lower() not in SUPPORTED_PROVIDERS:
        raise HTTPException(status_code=400, detail=f"Unknown provider: {provider}")
    
    deleted = delete_key(API_KEY, provider)
    if deleted:
        return {"ok": True, "message": f"{provider} key deleted"}
    else:
        return {"ok": False, "message": f"No key found for {provider}"}


# -----------------------------
# Routes: Chat
# -----------------------------

# Privacy & Security: Request size limits
MAX_MESSAGE_LENGTH = 10_000_000  # ~10MB - aligned with largest LLM context windows (Gemini 2M tokens)


@app.post("/runs", response_model=RunCreateResponse)
async def create_run(
    message: str = Form(...),
    configs: str = Form(...),  # JSON string
    max_parallel: int = Form(10),
    attachments: List[UploadFile] = File(default=[]),
    user: Optional[UserInfo] = Depends(require_api_key_or_jwt)
) -> RunCreateResponse:
    # Input validation: message length
    if len(message) > MAX_MESSAGE_LENGTH:
        raise HTTPException(status_code=400, detail=f"Message exceeds {MAX_MESSAGE_LENGTH} character limit")
    
    
    # Parse configs JSON
    try:
        config_list = json.loads(configs)
        config_objs = [ModelConfig(**cfg) for cfg in config_list]
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid configs JSON: {str(e)}")

    if not config_objs or len(config_objs) > 10:
        raise HTTPException(status_code=400, detail="configs must be 1..10 items")

    # -----------------------------
    # Quota & Platform Key Handling
    # -----------------------------
    # Determine which bots use BYOK vs platform keys
    # - BYOK: User provides their own API key (no quota usage)
    # - Platform: Uses botchat's API key (counts against quota)
    platform_key_config_ids: List[str] = []
    
    for cfg in config_objs:
        has_byok = bool(cfg.provider_key)  # User provided their own key
        
        if not has_byok:
            # Check if platform has a key for this provider
            platform_key = get_platform_key(cfg.provider)
            if platform_key:
                # Inject platform key into config
                cfg.provider_key = platform_key
                platform_key_config_ids.append(cfg.id)
            else:
                # No BYOK and no platform key - try server env fallback
                # This maintains backward compatibility with existing deployments
                pass
    
    # Check quota if using any platform keys
    if platform_key_config_ids and user:
        quota = await get_user_quota(user.provider, user.user_id, user.email)
        
        # Count how many messages will use platform keys
        platform_message_count = len(platform_key_config_ids)
        
        if quota["remaining"] < platform_message_count:
            # Quota exceeded
            if quota["remaining"] == 0:
                raise HTTPException(
                    status_code=429, 
                    detail=f"Message quota exhausted. Used {quota['used']}/{quota['limit']} messages this period. "
                           f"Upgrade to paid for {PAID_TIER_QUOTA} messages/month, or add your own API keys in Settings > Advanced."
                )
            else:
                # Partial quota - let user know some bots won't work
                raise HTTPException(
                    status_code=429,
                    detail=f"Insufficient quota. You have {quota['remaining']} messages remaining but selected {platform_message_count} bots using platform keys. "
                           f"Use fewer bots or add your own API keys in Settings > Advanced."
                )
    
    # If no user and platform keys needed, require authentication
    if platform_key_config_ids and not user:
        raise HTTPException(
            status_code=401,
            detail="Authentication required for platform API key usage. Sign in or provide your own API keys."
        )

    # MSGMODEL v3.2.0: In-Memory File Handling with Multi-File Support
    # Files are read into memory as bytes (not written to disk).
    # This ensures:
    # 1. Privacy: Files never touch the server filesystem
    # 2. Stateless: No cleanup required; each request is independent
    # 3. Smart: PDFs are intelligently handled per-provider
    #    - Gemini: Native PDF support (file passed as-is)
    #    - OpenAI: Text extraction (PDF converted to text)
    file_list: List[Tuple[bytes, str]] = []
    if attachments:
        for attachment in attachments:
            try:
                file_bytes = await attachment.read()
                filename = attachment.filename or "unknown"
                content_type = attachment.content_type or "unknown"
                
                file_list.append((file_bytes, filename))
                # INFO level log to diagnose mobile attachment issues
                logger.info("Attachment received: filename='%s', content_type='%s', size=%d bytes", 
                           filename, content_type, len(file_bytes))
                
                # Log PDF detection
                if filename.lower().endswith('.pdf'):
                    logger.info("PDF detected - will extract text for OpenAI, pass native PDF to Gemini")
            except Exception as e:
                logger.error("Failed to read attachment %s: %s", attachment.filename, str(e))
                raise HTTPException(status_code=400, detail=f"Failed to read attachment: {str(e)}")
    
    if not file_list:
        logger.debug("No attachments received")

    run_id = str(uuid.uuid4())
    run = RunState(
        run_id=run_id,
        created_at=time.time(),
        user_info=user,
        platform_key_configs=platform_key_config_ids,
    )

    async with RUNS_LOCK:
        RUNS[run_id] = run

    # Kick off background orchestration with file_list (will be processed per-provider)
    asyncio.create_task(run_fanout(run, config_objs, message, max_parallel, file_list if file_list else None))

    return RunCreateResponse(run_id=run_id)


@app.get("/runs/{run_id}/events")
async def stream_events(run_id: str, api_key=Depends(require_api_key)):
    async with RUNS_LOCK:
        run = RUNS.get(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="run not found")

    async def event_generator():
        try:
            # Initial hello (helps the UI know it connected)
            yield sse("hello", {"run_id": run_id})

            # Heartbeat so proxies don’t buffer/close
            last_ping = time.time()

            while True:
                # If done and queue empty, close stream
                if run.done and run.queue.empty():
                    yield sse("goodbye", {"run_id": run_id})
                    break

                try:
                    msg = await asyncio.wait_for(run.queue.get(), timeout=0.5)
                    yield msg
                except asyncio.TimeoutError:
                    pass

                if time.time() - last_ping > 10:
                    last_ping = time.time()
                    yield sse("ping", {"t": last_ping})
        finally:
            if run.done:
                async with RUNS_LOCK:
                    if RUNS.get(run_id) is run:
                        RUNS.pop(run_id, None)
                        logger.debug("Purged run %s after stream closed", run_id)

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.post("/runs/{run_id}/synthesize")
async def synthesize(run_id: str, req: SynthesizeRequest, api_key=Depends(require_api_key)):
    async with RUNS_LOCK:
        run = RUNS.get(run_id)
    if not run:
        raise HTTPException(status_code=404, detail="run not found")
    if not run.finals:
        raise HTTPException(status_code=400, detail="no panel finals yet")

    asyncio.create_task(fake_synthesize(run, req))
    return {"ok": True}